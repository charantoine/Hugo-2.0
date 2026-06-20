#!/usr/bin/env python3
"""Generate CTO briefing docx — Hugo 2.0 convergence test campaign."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(__file__).resolve().parents[1] / "briefing_cto_convergence_hugo_2_0_tests.docx"


def add_bullet(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def main() -> None:
    doc = Document()
    title = doc.add_heading("Briefing CTO — Convergence Hugo 2.0", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Campagne de tests complétude · {date.today().isoformat()}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Objet de ce document", level=1)
    doc.add_paragraph(
        "Ce briefing synthétise le travail de convergence réalisé localement entre le Hugo réel "
        "(backend Django + front apprenant 1.8) et la cible spec 2.0, ainsi qu’une campagne de "
        "tests backend et front exécutée le 18 juin 2026. Il vise à vous permettre de décider "
        "quoi conserver, quoi refaire, et comment mettre en ligne de manière sécurisée pour validation."
    )

    doc.add_heading("2. Ce qui a été fait (réel livré local)", level=1)
    doc.add_paragraph(
        "Lots documentés clusters 2 à 16 + livraison 20/06 profils conversationnels globaux apprenant. "
        "Matrice 13 domaines (V6). Côté produit :"
    )
    for item in [
        "UIState canonique (scène, progression, CTA, conversation_mode, profils d’affichage).",
        "Posture apprenant : API set-posture + PostureSelector (transitions backend).",
        "Mémoire intra-conversation : memory-summary + panneau LearnerMemoryPanel.",
        "CTA synthèse/évaluation dont état advisory « déconseillé mais possible ».",
        "Bandeau scène/objectif/maturité/dispersion (LearnerSceneContextBar).",
        "Trois profils d’affichage youth/adult/professional (même contenu, rendu différencié).",
        "Profil conversationnel global apprenant (diag + réflexif + bûchage + évaluation finale) — admin + résolution runtime.",
        "Affectation profil global au groupe ; fallback legacy TutorPrompt conservé.",
        "Interfaces formateur V0 : tableau knowledge, validation, élicitation.",
        "Garde-fous confidentialité : pas de P0/verbatim en UIState ; timeline tuteur filtrée.",
        "Multi-tenant : scope org, tutor-links SUPERADMIN-only, Playwright tenant 11/11 (20/06).",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3. Résultats campagnes de tests (local)", level=1)
    doc.add_paragraph("Environnement : local · Backend pytest · Front Playwright (Vite proxy /api).")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campagne"
    hdr[1].text = "Résultat"
    hdr[2].text = "Commentaire"
    rows = [
        ("Backend convergence 18/06 (168 tests)", "168 PASS", "Clusters 3/4/15/16, mémoire, CTA, RAG, observabilité"),
        ("Playwright 18/06 (20 tests)", "19 PASS / 1 FAIL", "FAIL = libellé titre formateur obsolète dans le test"),
        ("Cluster 16 apprenant (10 tests)", "10 PASS", "Posture, mémoire, scène, profils A1/A2/A3"),
        ("Multi-tenant smoke 20/06 relaxed", "90 PASS, 3 SKIP", "Isolation org, tutor-links, ACL encadrants"),
        ("Playwright tenant 20/06", "11 PASS", "SUPERADMIN switcher, ORGADMIN, TUTEUR, APPRENANT"),
        ("Profils globaux 20/06", "12 PASS dédiés", "CRUD, résolution, legacy template, priorité session/groupe"),
        ("Audit pipes convo 20/06", "26 + 32 PASS", "Pipeline admin↔runtime ; non-régression C15/C16"),
        ("Playwright admin profils 20/06", "4 PASS", "Création profil + affectation groupe SUPERADMIN"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val

    doc.add_heading("4. Ce qui fonctionne à coup sûr (preuves automatisées)", level=1)
    for item in [
        "Contrats UIState sans fuite P0 (INV-01, cluster 3/4/16).",
        "Mémoire gouvernée sans verbatim dans session_memory (domaine 20).",
        "CTA backend-driven cohérents avec blocking_reasons et advisory.",
        "Posture : transitions API + refus propre + sync UIState.",
        "Confidentialité multi-tenant : IDOR ui-state, verbatim masqué tuteur (B1-01).",
        "Exports ORGADMIN tenant-scoped ; trainer/tutor bloqués.",
        "RAG lexical gouverné (pas vectoriel).",
        "Interface apprenant E2E : posture, progression, mémoire panneau, 3 profils.",
        "Admin profils globaux : création, édition, affectation groupe (Playwright 20/06).",
        "Résolution runtime profil global → slots diag/réflexif/bûchage/éval (pytest).",
        "Tuteur / COORDO / exports orgadmin Playwright OK ; tenant multi-org 11/11.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("5. Ce qui ne fonctionne pas ou reste non prouvé", level=1)
    doc.add_paragraph("Distinction : échec test vs lacune produit vs non testé en prod.")
    for item in [
        ("Échec Playwright trainer smoke (18/06)", " — titre attendu obsolète. API et tableau OK. Correction = test uniquement."),
        ("Encoors / prod distante", " — parité non tranchée (A_VÉRIFIER)."),
        ("RLS PostgreSQL prod / CI strict", " — gate OK ; connexion DB CI/local à aligner sur rôle non-superuser."),
        ("Qualité conversationnelle bout-en-bout LLM", " — P0/phases testés unitairement ; pas campagne multi-tours LLM réel."),
        ("Prompts synthèse admin", " — hardcodés synthesis_service.py ; aucune UI (P1)."),
        ("Paramètres orchestrateur statiques", " — lecture seule admin ; non éditables (P2)."),
        ("Starter prompts admin", " — API CRUD sans page admin (P2)."),
        ("Policies évaluation par groupe", " — backend OK ; UI org-level surtout (P1)."),
        ("session.learner_conversation_profile REST", " — modèle + résolveur ; pas d’API session (décision produit)."),
        ("Intercalaires v1 (domaine 120)", " — CIBLE, 0 implémentation."),
        ("RAG vectoriel", " — CIBLE, pgvector inactif."),
        ("Choix profil cohorte IFT-042", " — CIBLE, pas surface ORGADMIN dédiée."),
        ("Verrou posture « séance avancée »", " — PARTIEL ; transitions posture OK."),
        ("Injection mémoire gouvernée dans prompt tour", " — hors lot 1, non livré."),
        ("Mémoire inter-sessions (LearnerThemeMemory)", " — stockage préparé, pas injecté runtime tour."),
    ]:
        add_bullet(doc, item[0], item[1] if len(item) > 1 else None)

    doc.add_heading("6. Prochaines développements recommandés (priorisation)", level=1)
    doc.add_paragraph("Ordre suggéré pour la vague 2.1 :")
    priorities = [
        ("P0 — Mise en ligne sécurisée", "Staging + smoke pytest + Playwright CI ; pas prod directe."),
        ("P0 — CI RLS", "Aligner TEST_DB_USER sur hugo_app ; corriger fixture strict (5 ERROR connus)."),
        ("P1 — Pipes conversationnels", "Admin synthèse ; UI policies évaluation groupe ; corriger smoke trainer."),
        ("P1 — Stabilisation tests", "U16-S2b advisory ; scénarios manuels S16 ; Playwright message tour."),
        ("P1 — Encoors", "Oracle avec credentials ; comparer JSON local."),
        ("P2 — Qualité conversationnelle", "Campagne multi-configurations postures × maturités × profils."),
        ("P2 — Starter prompts + orchestrateur", "Pages admin minimales ou décision reporter."),
        ("P3 — Couronne", "Intercalaires, RAG vectoriel, D9bis dashboards."),
    ]
    for label, desc in priorities:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(label + " : ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("7. Mise en ligne sécurisée — proposition", level=1)
    doc.add_paragraph(
        "Recommandation : déploiement par paliers sur environnement de préproduction dédié, "
        "sans exposition des surfaces tester/debug."
    )
    steps = [
        "Branche feature convergence — revue diff limitée aux domaines 10/20/31/70/110.",
        "Migration Django + bootstrap smoke (utilisateurs démo isolés).",
        "Pipeline CI : pytest noyau (168 tests) + Playwright cluster16 + smokes encadrants.",
        "Checklist manuelle persona A1/A2/A3 (30 min) — protocole cluster16 §4.",
        "Oracle Encoors read-only sur 1 session démo.",
        "Feature flags : désactiver v17 P0 prod tant qu’A_VÉRIFIER.",
        "Rollout apprenant d’abord ; formateur/tuteur ensuite ; ORGADMIN exports en dernier.",
    ]
    for s in steps:
        add_bullet(doc, s)

    doc.add_heading("8. Décision d’architecture — conserver / refaire / partir d’où", level=1)
    doc.add_paragraph(
        "Recommandation : conserver le socle backend-first actuel (UIState, CTA, mémoire API, "
        "progression contract) et le front ProdLearnerWorkspace enrichi clusters 15–16. "
        "Ne pas refondre le moteur P0 ni la spec 2.0."
    )
    decisions = [
        ("Conserver", "ui_state_builder, cta_ui_state, session_memory, orchestrateur existant, routes /app prod."),
        ("Conserver avec polish", "PostureSelector, LearnerMemoryPanel, profils CSS, bandeau scène, admin profils globaux."),
        ("Compléter (pas refaire)", "Admin synthèse, policies eval groupe, tests E2E advisory, IFT-042, doc SW-xx."),
        ("Reporter (couronne)", "Intercalaires, mémoire inter-sessions injection, RAG vectoriel, D9bis produit."),
        ("Refaire uniquement si audit prod diverge", "Encoors — comparer oracle avant merge prod."),
    ]
    for label, text in decisions:
        p = doc.add_paragraph()
        run = p.add_run(label + " : ")
        run.bold = True
        p.add_run(text)

    doc.add_heading("9. Documents de référence", level=1)
    refs = [
        "docs-workspace/README.md (§7 méthode tests, §7.6 pipes conversationnels)",
        "docs-workspace/test_final_complétude_améliorations.md (rapport 18/06 + addendum 20/06)",
        "docs-workspace/MINI_SPEC_PROFILS_CONVERSATIONNELS_APPRENANT.md",
        "docs-workspace/cluster2_matrice_runtime_vs_cible.md (V6)",
        "docs-workspace/cluster16_protocole_tests_interface_apprenant_v1.md",
        "docs-workspace/plan_tests_global_hugo.md",
        "docs-workspace/tests/archives/tests_hugo_2_0_2026-06-18_20.md",
        "docs-workspace/rapport_mise_a_jour_doc_post_cluster16_2026-06-18.md",
    ]
    for r in refs:
        add_bullet(doc, r)

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Document généré automatiquement — campagnes locales 2026-06-18 (convergence) et 2026-06-20 "
        "(multi-tenant + profils globaux apprenant). Encoors/prod restent A_VÉRIFIER jusqu’à oracle authentifié."
    )
    p.runs[0].italic = True

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    main()
